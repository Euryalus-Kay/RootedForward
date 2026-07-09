#!/usr/bin/env python3
# Build a formatted production-plan .docx for the Hyde Park film from the live
# data (data/hp-research.json + credits.json). Imports cleanly into Google Docs
# (upload to Drive, it opens as a formatted Doc; File > Download for PDF).
import json, os, re, html as _html
from docx import Document

def clean_credit(s):
    s = _html.unescape(s or "")
    s = re.sub(r"\s+", " ", s).strip(" .,;")
    return s
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
research = json.load(open(os.path.join(ROOT, "data/hp-research.json")))
credits = json.load(open(os.path.join(ROOT, "public/media/hyde-park/credits.json")))

FOREST = RGBColor(0x1E, 0x3A, 0x2B)
RUST = RGBColor(0xB0, 0x4A, 0x2E)
INK = RGBColor(0x22, 0x20, 0x1D)
GRAY = RGBColor(0x6B, 0x66, 0x5E)
SERIF, SANS = "Georgia", "Gill Sans"

ORDER = ["intro", "formation", "university", "worlds-fair", "color-line", "urban-renewal", "present"]
by = {c["chapterId"]: c for c in research["chapters"]}

PANOS = [
    ("pano-hyde-park", "57th Street Beach, facing the skyline across the lakefront"),
    ("pano-main-quad", "The center of the University of Chicago Main Quadrangles"),
    ("pano-jackson-park", "Jackson Park, by the Museum of Science and Industry and the Wooded Island"),
    ("pano-55th-street", "55th Street, or the University Park townhouses, the rebuilt ground"),
    ("pano-obama-center", "The Obama Presidential Center site in Jackson Park"),
]
BROLL = [("present-intro", "Chapter 01"), ("present-formation", "Chapter 02"), ("present-present", "Chapter 07")]

doc = Document()
# base style
st = doc.styles["Normal"]
st.font.name = SERIF; st.font.size = Pt(11); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.18

def shade(color="1E3A2B"):
    return None

def set_color(run, rgb): run.font.color.rgb = rgb

def eyebrow(text, color=RUST):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); r.font.name = SANS; r.font.size = Pt(9.5); r.bold = True
    r.font.color.rgb = color
    # letter spacing
    rPr = r._element.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), "40"); rPr.append(sp)
    return p

def heading(text, size=20, color=FOREST, space_before=14, sans=False):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = SANS if sans else SERIF; r.font.size = Pt(size); r.bold = True
    r.font.color.rgb = color
    return p

def body(text, italic=False, color=INK, size=11, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def quote(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11.5); r.italic = True; r.font.color.rgb = RGBColor(0x33, 0x30, 0x2B)
    # left rust bar via border
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18"); left.set(qn("w:space"), "10"); left.set(qn("w:color"), "B04A2E")
    pbdr.append(left); pPr.append(pbdr)
    return p

def bullet(text, sub=None):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(10.5)
    if sub:
        r2 = p.add_run("  " + sub); r2.font.size = Pt(10.5); r2.font.color.rgb = GRAY
    return p

def rule():
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "D8D1C5")
    pbdr.append(bot); pPr.append(pbdr)

# ---------------- title block ----------------
t = doc.add_paragraph(); t.paragraph_format.space_before = Pt(60); t.paragraph_format.space_after = Pt(2)
r = t.add_run("Hyde Park, Built and Rebuilt"); r.font.name = SERIF; r.font.size = Pt(34); r.bold = True; r.font.color.rgb = FOREST
s = doc.add_paragraph(); s.paragraph_format.space_after = Pt(20)
r = s.add_run("Production plan, script, and shot list"); r.font.name = SERIF; r.italic = True; r.font.size = Pt(16); r.font.color.rgb = RUST
eyebrow("A history of the ground beneath a Chicago neighborhood  .  Rooted Forward")
body("An immersive documentary for rooted-forward.org. Seven chapters carry Hyde Park from Paul Cornell's "
     "1853 lakefront bet through the University of Chicago, the 1893 World's Fair, the racial covenants, "
     "university-led urban renewal, and the present-day Obama Center. The cut already plays end to end with "
     "labeled placeholders, so you can watch it before filming anything.", color=GRAY)

# ---------------- how it works ----------------
heading("How this works", 16)
body("The archival photographs in the cut are real, public-domain or Creative Commons images, and they stay. "
     "The narration you hear now is a scratch machine voice. You re-record every voiceover below in your own "
     "voice, then film the host pieces, the 360 look-arounds, and the present-day b-roll and swap them in for "
     "the placeholders. Upload a clip in /admin/studio or /admin/immersive and point the matching id at it; "
     "the cut updates in place.")

heading("Everything you film, in one place", 15, color=RUST, sans=True)
eyebrow("Your two pieces to camera")
bullet("host-intro", "Chapter 01 opening, on a Hyde Park sidewalk")
bullet("host-close", "Chapter 07 close, in Jackson Park by the Obama Center")
eyebrow("360 captures with your 3D camera")
for pid, desc in PANOS: bullet(pid, desc)
eyebrow("Present-day b-roll")
for bid, ch in BROLL: bullet(bid, ch)

doc.add_page_break()

# ---------------- chapters ----------------
heading("The chapters", 22)
num = 0
for cid in ORDER:
    c = by[cid]; num += 1
    title = c.get("working") or cid
    era = c.get("era", "")
    rule()
    eyebrow(f"Chapter {num:02d}  .  {era}")
    heading(title, 18, space_before=2)
    sc = c.get("script", {})
    host = (sc.get("hostScript") or "").strip()
    if host:
        eyebrow("Say this on camera", color=FOREST)
        quote(host)
    vo = (sc.get("voiceover") or "").strip()
    if vo:
        eyebrow("Record this voiceover", color=FOREST)
        quote(vo)
    shot = (sc.get("shotNotes") or "").strip()
    if shot:
        eyebrow("Shots and direction", color=FOREST)
        body(shot)
    # archival stills for this chapter
    ims = [(k, v) for k, v in credits.items() if k.startswith(cid + "-")]
    if ims:
        eyebrow("Archival stills already in the cut", color=FOREST)
        for k, v in sorted(ims):
            name = clean_credit(re.sub(r"\.(jpg|png|tif|JPG)$", "", v.get("commonsTitle", k).replace("File:", "")))
            lic = clean_credit(v.get("license", ""))
            who = clean_credit(v.get("artist", ""))
            bullet(name, lic + (f"  .  {who}" if who and who.lower() != "unknown" else ""))

# ---------------- whats built ----------------
doc.add_page_break()
heading("What the film already does", 22)
body("This is not a storyboard on paper. The cut is rendered and runs about six and a half minutes. Built into it:")
for h, t in [
    ("A warm-duotone archival look", "Every historical photo is graded to a consistent sepia and shown complete over a blurred copy of itself, so faces and buildings are never cropped."),
    ("An animated locator map and master timeline", "A stylized South Side map shows Washington Park and Jackson Park framing Hyde Park, and a 1853-to-today timeline runs as a hero sequence and as a lit ribbon that carries across every chapter."),
    ("Two verified data charts", "The renewal return-rate gap as paired bars, 46 percent of white families stayed versus 17 percent of Black families, and the Black population growth from about 40,000 in 1910 to about 278,000 by 1940 as an animated census curve."),
    ("Museum wall labels", "Fifteen labels emerge on the beat the narration names a subject (the Ferris Wheel, the Court of Honor, Rockefeller, the Hansberry House), each pointing at a confirmed detail."),
    ("Dimensional stat moments", "The hero numbers (300 acres, 594 students, about 4,000 families) are extruded into depth and count up, then hold."),
    ("Disciplined motion", "Each shot makes one slow move and settles, sized to read on a large wall."),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(h + ".  "); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = FOREST
    r2 = p.add_run(t); r2.font.size = Pt(11)

heading("Ground rules the film keeps", 15, color=RUST, sans=True)
for t in [
    "Real, verified facts only. No invented numbers, quotes, dates, or sources.",
    "Every image is genuinely public domain or Creative Commons, credited on screen and in the sources.",
    "No em-dashes and no colons inside sentences in any on-screen text, the most reliable AI tells.",
]:
    bullet(t)

# ---------------- sources ----------------
doc.add_page_break()
heading("Verified sources", 22)
body("Pulled from the fact-check pass. Each chapter's claims were checked against these.", color=GRAY)
seen = set()
for cid in ORDER:
    c = by[cid]
    srcs = (c.get("verify", {}) or {}).get("sources") or (c.get("research", {}) or {}).get("sources") or []
    srcs = [s for s in srcs if s and s not in seen]
    if not srcs: continue
    for s in srcs: seen.add(s)
    eyebrow(c.get("working", cid), color=FOREST)
    for s in srcs:
        bullet(s if isinstance(s, str) else json.dumps(s))

out_dir = os.path.join(ROOT, "docs")
out = os.path.join(out_dir, "Hyde-Park-Production-Plan.docx")
doc.save(out)
print("WROTE", out, "(", num, "chapters,", len(seen), "sources )")
